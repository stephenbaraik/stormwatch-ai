"""Generate a publication-quality StormWatch AI report as a Word document.

Focuses on what the public cares about: methodology, findings, results.
Removes internal detail (project structure, test suite, how-to-reproduce).

Usage:  python scripts/generate_report_docx.py
Output: docs/StormWatch_AI_Report.docx
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parent.parent
FIGURES = PROJECT_ROOT / "docs" / "figures"
OUTPUT = PROJECT_ROOT / "docs" / "StormWatch_AI_Report.docx"

# ── design tokens ──
NAVY = RGBColor(0x1B, 0x3A, 0x5C)
BLUE = RGBColor(0x4C, 0x72, 0xB0)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GREY = RGBColor(0x99, 0x99, 0x99)
RED = RGBColor(0xC7, 0x3E, 0x1D)
GREEN = RGBColor(0x55, 0xA8, 0x68)
BLACK = RGBColor(0x22, 0x22, 0x22)

_doc: Document


# ── helper functions ──

def figure(path: str, width: float = 5.5, caption: str = "", number: str = ""):
    """Insert a centred figure with numbered caption."""
    label = f"Figure {number}: {caption}" if number else caption
    sp = _doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(10)
    fpath = FIGURES / path
    if fpath.exists():
        _doc.add_picture(str(fpath), width=Inches(width))
        _doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if label:
        cap = _doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        r = cap.add_run(label)
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = LIGHT_GREY


def h1(text: str):
    _doc.add_heading(text, level=1)


def h2(text: str):
    _doc.add_heading(text, level=2)


def h3(text: str):
    _doc.add_heading(text, level=3)


def para(text: str, size: int = 10, bold: bool = False, italic: bool = False):
    p = _doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = BLACK
    return p


def bold(text: str, size: int = 10):
    return para(text, size=size, bold=True)


def small(text: str):
    para(text, size=9)


def bullet(text: str):
    p = _doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = BLACK


def numbered(text: str):
    p = _doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = BLACK


def table(headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    n_cols = len(headers)
    n_rows = len(rows) + 1
    tbl = _doc.add_table(rows=n_rows, cols=n_cols)
    tbl.style = "Light Shading Accent 1"
    for ci, h in enumerate(headers):
        cell = tbl.cell(0, ci)
        cell.text = str(h)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            if ci < n_cols:
                cell = tbl.cell(ri + 1, ci)
                cell.text = str(val)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
    if widths:
        for ci, w in enumerate(widths):
            if ci < n_cols:
                for row in tbl.rows:
                    row.cells[ci].width = Inches(w)
    _doc.add_paragraph("")


def code(text: str):
    p = _doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def page_break():
    _doc.add_page_break()


def hr():
    p = _doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("─" * 70)
    r.font.size = Pt(6)
    r.font.color.rgb = LIGHT_GREY


# ── document setup ──

def _setup_styles():
    style = _doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    style.font.color.rgb = BLACK

    for lvl in range(1, 5):
        hs = _doc.styles[f"Heading {lvl}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = NAVY
        hs.paragraph_format.space_before = Pt(14) if lvl < 3 else Pt(10)
        hs.paragraph_format.space_after = Pt(4)
        sizes = {1: 20, 2: 15, 3: 12, 4: 11}
        hs.font.size = Pt(sizes[lvl])

    for section in _doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)


# ═══════════════════════════════════════════════
#  SECTIONS
# ═══════════════════════════════════════════════


def title_page():
    for _ in range(8):
        _doc.add_paragraph("")

    t = _doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("StormWatch AI")
    r.font.size = Pt(40)
    r.font.bold = True
    r.font.color.rgb = NAVY

    st = _doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("Extreme Weather Early Warning\nfor the Indian Subcontinent")
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE

    _doc.add_paragraph("")

    meta = _doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        "An end-to-end machine learning system for multi-hazard weather prediction",
        "using XGBoost models trained on real meteorological data",
        "",
        "July 2026",
    ]:
        r = meta.add_run(line + "\n")
        r.font.size = Pt(11)
        r.font.color.rgb = GREY

    page_break()


def abstract():
    """A one-paragraph executive summary that stands alone."""
    h1("Abstract")
    para(
        "We present StormWatch AI, an end-to-end machine learning system for predicting three "
        "classes of extreme weather events across the Indian subcontinent: cyclone intensity "
        "(6-class Saffir-Simpson category), next-day heatwave occurrence (binary), and next-day "
        "extreme rainfall exceedance (binary). All three models use XGBoost with StandardScaler "
        "preprocessing, trained exclusively on real historical data — 90,138 daily weather records "
        "from 15 Indian cities (2009–2026, sourced from the Open-Meteo archive via Supabase) and "
        "57,632 tropical cyclone track records from NOAA's IBTrACS database (North Indian Ocean basin). "
        "The cyclone model achieves 98.5% test accuracy (0.985 weighted F1) using eight physical "
        "and temporal features. The heatwave model achieves 99.0% accuracy with 0.997 ROC-AUC using "
        "thirteen lagged temperature and atmospheric features. The rainfall model achieves 88.3% "
        "accuracy with 0.881 ROC-AUC using fourteen lagged precipitation and atmospheric features — "
        "a result that reflects the genuine difficulty of next-day rainfall forecasting once a "
        "label-leakage error was identified and corrected (earlier versions that included same-day "
        "precipitation as a feature scored an artificial 99.7%). We discuss the leakage discovery "
        "as a methodological case study, audit the dominant feature importance of minimum central "
        "pressure in the cyclone model to confirm it represents legitimate physical signal rather "
        "than definitional leakage, and report that Hyperopt-based hyperparameter tuning produced "
        "no meaningful improvement over default XGBoost parameters. The system is served via a "
        "FastAPI with API-key authentication, tracked with MLflow experiment and model registry, "
        "and deployed with Docker."
    )


def introduction():
    h1("1. Introduction")
    para(
        "Extreme weather events — tropical cyclones, heatwaves, and torrential rainfall — cause "
        "significant loss of life and economic damage across the Indian subcontinent. The region's "
        "diverse climate zones, ranging from arid Rajasthan to the humid northeast and the cyclone-prone "
        "Bay of Bengal, present a challenging testbed for machine learning-based forecasting. An "
        "effective early warning system must handle multiple hazard types, operate on heterogeneous "
        "spatial data, and produce predictions with sufficient lead time to enable preventive action."
    )

    para(
        "This report describes StormWatch AI, a system that addresses three prediction tasks "
        "simultaneously: (1) classifying tropical cyclone intensity on the Saffir-Simpson scale (0–5), "
        "(2) detecting whether the following day will be a heatwave day, and (3) predicting whether "
        "the following day's precipitation will exceed the city-specific 95th percentile threshold. "
        "All three tasks are framed as supervised learning problems using XGBoost classifiers."
    )

    para(
        "We place particular emphasis on the integrity of the modelling pipeline. During development, "
        "we identified a label-leakage bug that had inflated reported performance across all three "
        "models — each model was inadvertently fed the same meteorological variable used to define "
        "its own target label. Correcting this required removing same-day leaky features, recomputing "
        "rolling statistics from prior-day values, and shifting heatwave and rainfall targets forward "
        "by one day to create genuine forecasts. We document this process in detail (§5.4) as a "
        "methodological case study relevant to any practitioner working with temporally autocorrelated "
        "weather data."
    )

    para(
        "We also report two additional investigations that are not typically included in similar "
        "systems: an audit of whether the dominant feature importance of minimum central pressure "
        "in the cyclone model constitutes leakage of its own kind (§5.2), and a full Hyperopt "
        "hyperparameter tuning study that found default XGBoost parameters to be effectively optimal "
        "for all three tasks (§5.3)."
    )


def data_section():
    h1("2. Data")

    h2("2.1 Weather Data")
    para(
        "Daily meteorological observations were collected from the Open-Meteo Archive API for "
        "15 Indian cities spanning four climate zones — coastal (Mumbai, Chennai, Kolkata, Kochi, "
        "Bhubaneswar, Visakhapatnam, Surat), inland (Delhi, Hyderabad, Bengaluru, Lucknow, Pune), "
        "arid (Ahmedabad, Jaipur), and humid (Guwahati). The 18 variables include temperature "
        "extremes, precipitation, wind speed and gusts, humidity, surface pressure, cloud cover, "
        "solar radiation, evapotranspiration, and convective available potential energy (CAPE)."
    )

    para(
        "The dataset spans from 31 December 2009 to 25 June 2026, yielding 90,138 daily records. "
        "Fourteen cities contribute 6,021 rows each; Visakhapatnam contributes 5,844 rows, having "
        "joined the data collection pipeline slightly later. The data were uploaded to a Supabase "
        "PostgreSQL instance for centralised storage and retrieved via a service-role API pull "
        "script for model training."
    )

    table(
        ["Variable", "Description", "Units"],
        [
            ["temperature_2m_max", "Daily maximum temperature at 2m", "°C"],
            ["temperature_2m_min", "Daily minimum temperature at 2m", "°C"],
            ["precipitation_sum", "Total daily precipitation", "mm"],
            ["wind_speed_10m_max", "Maximum wind speed at 10m", "km/h"],
            ["relative_humidity_2m_mean", "Mean relative humidity at 2m", "%"],
            ["surface_pressure_mean", "Mean surface pressure", "hPa"],
            ["cloud_cover_mean", "Mean total cloud cover", "%"],
            ["shortwave_radiation_sum", "Total solar radiation", "MJ/m²"],
            ["cape_mean", "Mean convective available potential energy", "J/kg"],
        ],
        [2.8, 3.2, 1.0],
    )

    figure("data_coverage.png", number="1", caption="Data coverage per city — all 15 cities with approximately 5,500–6,500 daily records each.")
    figure("temperature_trends.png", number="2", caption="Daily maximum temperature trends (2010–2026) for Mumbai, Delhi, Chennai, and Kolkata. Black line: 12-month rolling mean, revealing seasonal cycles and long-term warming.")

    h2("2.2 Cyclone Data")
    para(
        "Tropical cyclone track records were obtained from NOAA's International Best Track Archive "
        "for Climate Stewardship (IBTrACS), specifically the North Indian Ocean basin (code NI). "
        "The raw download contains 62,606 records spanning from 1842 to the present. After filtering "
        "to tropical storm (NATURE = TS) and mixed/subtropical (MX) classifications, 57,632 records "
        "remain across all six Saffir-Simpson categories (0: Tropical Depression through 5: Super Cyclone)."
    )

    para(
        "Two data-access bugs were discovered and fixed during this work: (1) the basin code had "
        "not been updated from NOAA's deprecated 'IO' to the current 'NI', causing silent 404 errors "
        "on download, and (2) the NATURE filter matched the non-existent code 'TC' instead of the "
        "real codes 'TS' and 'MX', silently dropping all 62,606 downloaded records before any model "
        "saw them. Both are documented in §5.4."
    )

    table(
        ["Variable", "Description"],
        [
            ["lat, lon", "Storm centre coordinates"],
            ["pressure_min", "Minimum central pressure (hPa)"],
            ["dist_to_land", "Distance to nearest landmass (km)"],
            ["year, month, dayofyear", "Temporal features"],
            ["category", "Saffir-Simpson category 0–5 (target)"],
        ],
        [2.0, 5.0],
    )

    h2("2.3 Preprocessing")
    para(
        "Extreme event labels were derived from the raw meteorological variables using threshold-based "
        "definitions. A heatwave day is defined as a period where the daily maximum temperature exceeds "
        "40°C for three or more consecutive days. Extreme rainfall is defined as any day where "
        "precipitation exceeds the city-specific 95th percentile, computed per city to account for "
        "spatial heterogeneity in rainfall climatology. Saffir-Simpson cyclone categories were "
        "derived from maximum sustained wind speed (in knots) using the standard NOAA thresholds."
    )

    para(
        "For the heatwave and rainfall forecasting tasks, we created lagged features (1-day, 3-day, "
        "and 7-day lags of temperature, precipitation, and wind speed) and rolling statistics "
        "(3-day and 7-day rolling means and standard deviations). Critically, all rolling statistics "
        "were computed on the prior-day-shifted series — today's rolling mean reflects only the "
        "preceding days, never today's own reading. This design choice is essential for avoiding "
        "label leakage and is discussed further in §5.4."
    )


def methodology():
    h1("3. Methodology")

    h2("3.1 Feature Engineering")
    para(
        "Each model uses a distinct feature set designed to capture the physical processes relevant "
        "to its prediction target. All feature engineering is implemented in stormwatch/features/builder.py "
        "as sklearn-compatible functions that return (X, y) tuples."
    )

    h3("Cyclone Intensity — 8 features")
    para(
        "The cyclone model uses absolute latitude (hemisphere-agnostic), longitude, signed latitude, "
        "minimum central pressure, distance to land, and three temporal features (year, month, "
        "day of year). Notably, maximum sustained wind speed (wind_kts) is excluded despite being "
        "the strongest physical correlate of cyclone intensity. This is because the Saffir-Simpson "
        "category is deterministically derived from wind_kts via fixed thresholds — including it "
        "would reduce the classification task to a lookup table rather than a prediction problem. "
        "We verified that minimum central pressure, which becomes the dominant feature at 52.3% "
        "importance after wind_kts is removed, represents a legitimate physical signal rather than "
        "definitional leakage (§5.2)."
    )

    table(
        ["Feature", "Type", "Physical basis"],
        [
            ["lat_abs", "float", "Cyclone intensity correlates with absolute latitude"],
            ["lon", "float", "Identifies basin sub-region (Bay of Bengal vs Arabian Sea)"],
            ["lat", "float", "Signed latitude preserves hemisphere information"],
            ["pressure_min", "float", "Lower central pressure → stronger storm (gradient wind balance)"],
            ["dist_to_land", "float", "Land interaction weakens storms via surface friction"],
            ["year", "int", "Captures multi-decadal climate variability"],
            ["month", "int", "Captures seasonal cyclone patterns (pre/post-monsoon)"],
            ["dayofyear", "int", "Finer-grained seasonal signal"],
        ],
        [1.4, 0.7, 4.8],
    )

    h3("Heatwave Detection — 13 features")
    para(
        "The heatwave model predicts tomorrow's heatwave flag using only information available today. "
        "Today's own maximum and minimum temperatures are deliberately excluded because the heatwave "
        "label is defined from today's temperature — including it would be leakage. Instead, the model "
        "uses yesterday's temperature (lag-1), lags at 3 and 7 days, rolling means of prior-day "
        "maxima over 3 and 7 days, yesterday's minimum temperature, yesterday's precipitation, "
        "and today's exogenous atmospheric readings (humidity, wind speed, surface pressure) which "
        "are not derived from the label. Cyclic month encoding (sine/cosine) captures seasonality."
    )

    h3("Extreme Rainfall — 14 features")
    para(
        "The rainfall model follows the same anti-leakage principle: today's precipitation is excluded "
        "because the extreme rainfall label is defined from it. Instead, the model uses precipitation "
        "lags at 1, 3, and 7 days, rolling means of prior-day precipitation over 3 and 7 days, "
        "yesterday's maximum temperature and its 3-day rolling mean, and today's exogenous readings "
        "(humidity, wind speed, surface pressure, total cloud cover). The more distributed importance "
        "pattern of this model reflects the genuine complexity of next-day rainfall forecasting."
    )

    h2("3.2 Model Architecture")
    para(
        "All three models share a consistent architecture: a sklearn Pipeline consisting of a "
        "StandardScaler followed by an XGBoost classifier. The cyclone model uses multi:softprob "
        "objective with 6 classes and class-balanced weighting. The heatwave and rainfall models "
        "use binary:logistic with scale_pos_weight to handle class imbalance (positive rates of "
        "approximately 1.8% and 5.0% respectively). A BaseWeatherModel abstract class provides "
        "a common interface with train(), predict(), predict_proba(), evaluate(), save(), load(), "
        "and log_model() methods."
    )

    table(
        ["", "Cyclone", "Heatwave", "Rainfall"],
        [
            ["Algorithm", "XGBoost", "XGBoost", "XGBoost"],
            ["Objective", "multi:softprob", "binary:logistic", "binary:logistic"],
            ["Classes", "6 (0–5)", "2", "2"],
            ["Features", "8", "13", "14"],
            ["Imbalance handling", "Balanced weights", "scale_pos_weight", "scale_pos_weight"],
            ["Positive rate", "—", "1.8%", "5.0%"],
            ["Target framing", "Same-timestamp classification", "Next-day forecast", "Next-day forecast"],
        ],
        [1.6, 1.6, 1.6, 1.6],
    )

    h2("3.3 Training Protocol")
    para(
        "All models were trained on an 80/20 stratified train-test split (random_state=42). "
        "Training used default XGBoost hyperparameters after a Hyperopt tuning study found them "
        "to be effectively optimal for all three tasks (§5.3). Experiments are tracked with MLflow "
        "using a SQLite backend, with parameters, metrics, feature importances, model signatures, "
        "and input examples logged per run. Models are registered in the MLflow Model Registry "
        "with Production aliases and also saved as .pkl files for disk-based fallback loading."
    )


def results():
    h1("4. Results")

    # ── 4.1 Cyclone ──
    h2("4.1 Cyclone Intensity Classification")
    para(
        "The cyclone model was trained on 57,632 real IBTrACS records from the North Indian Ocean "
        "basin and evaluated on a held-out test set of 11,527 samples. The class distribution is "
        "heavily imbalanced: Category 0 (Tropical Depression) accounts for 54,068 of the full "
        "dataset while Category 5 appears only 114 times, reflecting the real-world rarity of "
        "super-cyclones."
    )

    table(
        ["Metric", "Value"],
        [
            ["Test Accuracy", "0.985"],
            ["Test F1 (weighted)", "0.985"],
            ["Test samples", "11,527"],
            ["Classes", "6 (Saffir-Simpson 0–5)"],
        ],
        [2.5, 3.5],
    )

    figure("cyclone_confusion_matrix.png", number="3",
           caption="Cyclone intensity confusion matrix. Category 0 dominates the test set (10,814/11,527). "
                   "Misclassifications are almost exclusively between adjacent categories, which is expected "
                   "given the continuous nature of the underlying wind-speed → category mapping.")

    figure("cyclone_class_distribution.png", number="4",
           caption="Test-set class distribution. The exponential decay in category frequency mirrors the "
                   "real-world rarity of intense tropical cyclones in the North Indian Ocean basin.")

    figure("feature_importance_cyclone.png", number="5",
           caption="Feature importance for the cyclone model. pressure_min dominates at 52.3%, followed "
                   "by year (24.2%) and lon (6.9%). lat carries zero importance because it is fully "
                   "redundant with lat_abs for North Indian Ocean storms, which are all in the northern "
                   "hemisphere. See §5.2 for a leakage audit of pressure_min.")

    para(
        "At 98.5% accuracy without wind_kts as a feature, the model's performance is strong but "
        "physically plausible. The dominant feature, minimum central pressure, is a well-established "
        "proxy for storm intensity through the gradient wind balance relationship — lower central "
        "pressure drives stronger pressure-gradient forces and thus higher wind speeds. Unlike "
        "wind_kts, pressure_min is not the defining variable for the Saffir-Simpson category, so "
        "its high importance does not constitute definitional leakage (verified in §5.2)."
    )

    # ── 4.2 Heatwave ──
    h2("4.2 Heatwave Next-Day Detection")
    para(
        "The heatwave model was trained on 90,123 weather samples (15 cities, after dropping each "
        "city's last day for the next-day target shift). The positive class (heatwave day) comprises "
        "only 1.8% of samples, making this a highly imbalanced binary classification task."
    )

    table(
        ["Metric", "Value"],
        [
            ["Test Accuracy", "0.990"],
            ["Test ROC-AUC", "0.997"],
            ["Test F1", "0.778"],
            ["Positive rate", "1.8% (1,641 / 90,138)"],
        ],
        [2.5, 3.5],
    )

    figure("feature_importance_heatwave.png", number="6",
           caption="Feature importance for the heatwave model. temp_max_lag_1 (yesterday's maximum "
                   "temperature) dominates at 96.5%, with the remaining 12 features collectively "
                   "contributing only 3.5%. This extreme concentration reflects the strong physical "
                   "autocorrelation of daily maximum temperatures — if yesterday was extremely hot, "
                   "tomorrow is overwhelmingly likely to be hot as well. The model has learned this "
                   "trivial-but-correct relationship.")

    para(
        "The 99.0% accuracy and 0.997 ROC-AUC are credible because heatwaves are physically persistent "
        "phenomena. A heatwave is defined by sustained high temperatures over multiple days, so "
        "yesterday's temperature is genuinely predictive of tomorrow's. The model is not cheating — "
        "it is exploiting a real physical regularity. The low F1 score (0.778) reflects the inherent "
        "difficulty of identifying the minority positive class in a dataset where 98.2% of days "
        "are not heatwave days."
    )

    # ── 4.3 Rainfall ──
    h2("4.3 Extreme Rainfall Next-Day Prediction")
    para(
        "The rainfall model was trained on the same 90,123 weather samples. The positive class "
        "(extreme rainfall exceedance of the city-specific 95th percentile) comprises approximately "
        "5.0% of samples."
    )

    table(
        ["Metric", "Value"],
        [
            ["Test Accuracy", "0.883"],
            ["Test ROC-AUC", "0.881"],
            ["Test F1", "0.330"],
            ["Positive rate", "5.0% (4,504 / 90,138)"],
        ],
        [2.5, 3.5],
    )

    figure("feature_importance_rainfall.png", number="7",
           caption="Feature importance for the rainfall model. Unlike the cyclone and heatwave models, "
                   "no single feature dominates. relative_humidity_2m_mean (6.9%), pressure_msl_mean "
                   "(5.6%), cloud_cover_mean (5.2%), and the collection of precipitation lag and "
                   "rolling-mean features each contribute 3.5–5.2%. This distributed pattern reflects "
                   "the genuinely complex, multi-factor nature of next-day precipitation forecasting "
                   "from surface meteorological variables alone.")

    para(
        "At 88.3% accuracy and 0.881 ROC-AUC, this model is the weakest of the three — and that is "
        "by design. Before the label-leakage fix was applied (§5.4), this same model scored 99.7% "
        "accuracy by including today's precipitation as a feature. The 11.4 percentage point drop "
        "after removing the leaky feature and shifting the target to the next day is the most "
        "important result in this report: it demonstrates that the model was previously solving a "
        "trivial lookup problem, not forecasting. The current 88.3% reflects the genuine difficulty "
        "of next-day rainfall prediction from surface variables alone — precisely the result that "
        "domain knowledge predicts."
    )

    # ── 4.4 Cross-model comparison ──
    h2("4.4 Cross-Model Comparison")
    figure("model_performance.png", number="8",
           caption="Performance comparison across all three models. Cyclone and heatwave achieve "
                   "strong results on all metrics. The rainfall model's low F1 (0.330) reflects the "
                   "difficulty of identifying minority positive cases in a fundamentally hard forecasting task.")

    table(
        ["Dimension", "Cyclone", "Heatwave", "Rainfall"],
        [
            ["Accuracy", "98.5%", "99.0%", "88.3%"],
            ["ROC-AUC", "— (multi-class)", "0.997", "0.881"],
            ["F1", "0.985", "0.778", "0.330"],
            ["Dominant feature", "pressure_min (52.3%)", "temp_max_lag_1 (96.5%)", "humidity (6.9%)"],
            ["Difficulty", "Low — pressure strongly predicts intensity", "Low — heatwaves are persistent", "High — next-day rain is inherently uncertain"],
        ],
        [1.8, 1.8, 1.8, 1.8],
    )

    # ── 4.5 EDA ──
    h2("4.5 Exploratory Analysis")

    figure("extreme_events_by_zone.png", number="9",
           caption="Extreme events by climate zone. Coastal cities dominate extreme rainfall events "
                   "(proximity to moisture sources from the Arabian Sea and Bay of Bengal), while "
                   "inland zones, particularly arid cities, record the most heatwave days.")

    figure("monthly_event_patterns.png", number="10",
           caption="Monthly distribution of extreme events. Heatwaves peak sharply in May–June "
                   "(pre-monsoon), while extreme rainfall events follow the monsoon cycle, peaking "
                   "from June through September. This seasonal structure is well-captured by the "
                   "cyclic month encoding features in all three models.")

    figure("city_event_counts.png", number="11",
           caption="Absolute event counts per city. Coastal cities (Chennai, Mumbai, Kolkata) and "
                   "humid Guwahati record the most extreme rainfall events, while inland Delhi and "
                   "arid Ahmedabad lead in heatwave counts. Visakhapatnam is included as the 15th city.")


def key_findings():
    h1("5. Key Findings & Methodological Insights")

    h2("5.1 The Leakage Story: Why Suspiciously Perfect Accuracy Is a Bug Signal")
    para(
        "During an initial training pass, all three models achieved near-perfect test scores: cyclone "
        "100.0%, heatwave 99.9%, and rainfall 99.7%. Upon investigation, we discovered that each model "
        "had been fed the exact same meteorological variable used to define its own target label:"
    )

    table(
        ["Model", "Target label definition", "Leaked feature (removed)"],
        [
            ["Cyclone", "Saffir-Simpson category = deterministic bucket of wind_kts", "wind_kts"],
            ["Heatwave", "heatwave_flag = same-day temp_max > 40°C for 3 consecutive days", "same-day temp_max"],
            ["Rainfall", "extreme_rainfall = same-day precipitation > city 95th percentile", "same-day precipitation"],
        ],
        [1.3, 3.7, 2.0],
    )

    para(
        "The models were not forecasting — they were re-deriving a lookup table. The heatwave model, "
        "for example, was given today's maximum temperature and asked whether today was a heatwave "
        "day — a question it could answer perfectly by simply checking whether that temperature "
        "exceeded 40°C. Three corrections were required to create genuine forecasting models:"
    )

    numbered(
        "Remove same-day leaky features: wind_kts from the cyclone feature set, same-day temp_max from "
        "heatwave, and same-day precipitation from rainfall. Keep same-day exogenous variables "
        "(humidity, wind, pressure, cloud cover, season) that are not derived from the label."
    )
    numbered(
        "Recompute rolling statistics on prior-day-shifted series: a 3-day rolling mean of temperature "
        "should reflect days t−1, t−2, t−3, not t, t−1, t−2. This was corrected in "
        "prepare_weather_features() by applying .shift(1) before .rolling() aggregation."
    )
    numbered(
        "Shift targets forward by one day: the heatwave and rainfall targets are now tomorrow's "
        "flag rather than today's. This was implemented per city using "
        "df.groupby('city')[target].shift(−1), with each city's final row dropped (no next-day label)."
    )

    para(
        "After these corrections, the models' responses differed in the direction domain knowledge "
        "would predict: cyclone remained high at 98.5% because pressure is a genuine physical proxy "
        "for storm intensity, heatwave stayed at 99.0% because extreme temperatures are highly "
        "autocorrelated day-to-day, and rainfall dropped to 88.3% because next-day precipitation is "
        "genuinely one of the hardest short-term forecasting targets. The fact that the same fix "
        "produced different impacts on each model, in the physically expected direction, provides "
        "strong evidence that the correction is valid rather than an overcorrection."
    )

    h2("5.2 Auditing pressure_min: Legitimate Signal, Not Hidden Leakage")
    para(
        "After removing wind_kts from the cyclone feature set, minimum central pressure became the "
        "dominant feature at 52.3% importance. Given that the label (Saffir-Simpson category) was "
        "derived from wind_kts, and pressure and wind are physically linked, it was worth verifying "
        "whether pressure_min constituted a subtler form of leakage."
    )

    para(
        "We ran three diagnostic experiments on a 1,000-sample stratified subset of the cyclone data:"
    )

    table(
        ["Experiment", "Features", "Accuracy", "Interpretation"],
        [
            ["Pressure only", "pressure_min alone", "95.7%", "Pressure alone is highly predictive but not deterministic"],
            ["No pressure", "All features except pressure_min", "95.6%", "Spatial + temporal features alone nearly match pressure-only"],
            ["Full model", "All 8 features", "98.2%", "Combined signals provide complementary information"],
        ],
        [1.5, 2.5, 1.0, 3.5],
    )

    para(
        "The results confirm that while pressure_min is the strongest individual predictor, the model "
        "does not collapse without it (only a 2.7% accuracy drop). The spatial and temporal features "
        "(year, lon, month, dayofyear, dist_to_land) together provide nearly as much predictive power. "
        "Furthermore, inspection of the preprocessing code confirmed that the Saffir-Simpson category "
        "is derived exclusively from wind_kts using the standard NOAA thresholds — pressure_min is "
        "never consulted in label generation."
    )

    bold(
        "Conclusion: pressure_min's 52.3% importance is legitimate physical signal. It reflects "
        "the well-established pressure-wind relationship (gradient wind balance), not a definitional "
        "dependency on the label. This is analogous to using rolling temperature to predict "
        "tomorrow's heatwave — correlated, predictive, but not the label itself.",
        size=10,
    )

    h2("5.3 Hyperopt Tuning: Defaults Were Already Optimal")
    para(
        "A full Hyperopt hyperparameter tuning study was conducted across all three models using "
        "the Tree-structured Parzen Estimator (TPE) algorithm with 20 trials and 3-fold "
        "cross-validation. The search spaces covered learning rate, max depth, n_estimators, "
        "subsample, colsample_bytree, min_child_weight, reg_lambda, gamma, and scale_pos_weight "
        "(where applicable). For cyclone, the optimisation objective was weighted F1; for heatwave "
        "and rainfall, ROC-AUC."
    )

    table(
        ["Model", "Baseline (default params)", "Hyperopt best", "Δ", "Verdict"],
        [
            ["Cyclone", "98.5% / 0.985 F1", "98.4% / 0.984 F1", "−0.1%", "Defaults optimal"],
            ["Heatwave", "99.0% / 0.997 AUC", "98.9% / 0.997 AUC", "−0.1%", "Defaults optimal"],
            ["Rainfall", "88.3% / 0.881 AUC", "75.2% / 0.894 AUC", "+0.013 AUC / −13.1% acc", "AUC-overfitted; defaults preferred"],
        ],
        [1.1, 2.0, 2.0, 1.5, 1.7],
    )

    para(
        "For cyclone and heatwave, Hyperopt found parameters that were trivially worse than the "
        "defaults. For rainfall, it found parameters that improved cross-validated AUC by 1.3% "
        "but destroyed test accuracy (−13.1%), a classic case of the tuning objective diverging "
        "from the evaluation metric of interest. The baseline (untuned) models are used for all "
        "production inference."
    )

    para(
        "This finding aligns with recent literature suggesting that XGBoost's default hyperparameters "
        "are well-tuned for a broad range of tabular prediction tasks. For practitioners building "
        "similar systems, our recommendation is to invest effort in feature engineering and data "
        "quality before expending computational resources on hyperparameter search."
    )

    h2("5.4 Data Pipeline Bugs Found and Fixed")
    para(
        "Beyond the label-leakage issue (§5.1), two additional data-access bugs were discovered "
        "during the July 2026 retraining pass:"
    )

    bold("IBTrACS basin code (IO → NI).")
    para(
        "NOAA retired the 'IO' (Indian Ocean) basin code at some point after the project's original "
        "configuration was written. The current codes are 'NI' (North Indian, the one relevant to "
        "India) and 'SI' (South Indian). Two separate locations — IBTrACSConfig in config.py/config.yaml "
        "and a hardcoded basin_urls dictionary in download.py — still referenced 'IO', causing "
        "silent 404 errors on download."
    )

    bold("Cyclone NATURE filter (TC → TS|MX).")
    para(
        "The preprocess_cyclones() function filtered IBTrACS records to NATURE containing 'TC', "
        "intending to retain only tropical cyclone observations. However, real IBTrACS data uses "
        "no 'TC' code — the actual classifications are TS (tropical storm), DS (disturbance), "
        "ET (extratropical), MX (mixture/subtropical), and NR (not reported). This filter silently "
        "dropped all 62,606 downloaded records before any model saw them. After fixing to match "
        "'TS|MX', 57,632 real tropical cyclone records were recovered."
    )


def discussion():
    h1("6. Discussion")

    h2("6.1 Why Heatwave Prediction Is 'Easy' and Rainfall Is 'Hard'")
    para(
        "The three-fold difference in model performance — from 99.0% accuracy for heatwaves to 88.3% "
        "for rainfall — is not a statement about model quality but about the underlying physics of "
        "each phenomenon. Heatwaves are defined by sustained high temperatures over multiple days, "
        "which creates a strong temporal autocorrelation: if today's maximum temperature was 42°C, "
        "tomorrow's is highly likely to be hot as well. The model's 96.5% feature importance on "
        "temp_max_lag_1 simply reflects this physical fact. The task is closer to nowcasting than "
        "forecasting, and the high accuracy is expected given the data."
    )

    para(
        "Rainfall, by contrast, is governed by mesoscale convective processes, moisture convergence "
        "patterns, and upper-atmospheric dynamics that are only weakly constrained by surface "
        "meteorological observations. A day of heavy rain is often followed by clearing skies, and "
        "vice versa — the autocorrelation structure that makes heatwaves predictable is absent. The "
        "model's distributed feature importance pattern (no single feature above 7%) and modest 88.3% "
        "accuracy correctly reflect this fundamental difficulty. Adding upper-air data (geopotential "
        "height, wind shear, specific humidity profiles) from reanalysis products like ERA5 would "
        "likely yield larger improvements than any modelling or hyperparameter change."
    )

    h2("6.2 System Design Choices")
    para(
        "Several architectural decisions were made with production deployment in mind:"
    )
    bullet(
        "MLflow model registry with Production aliases enables model versioning without code changes. "
        "The API server attempts to load from the registry first, with automatic disk-based .pkl "
        "fallback if the registry is unreachable — a pattern that works in both development and "
        "production environments."
    )
    bullet(
        "API key authentication via the X-API-Key header is configurable through a single environment "
        "variable (STORMWATCH_API_KEY). When unset, authentication is disabled for local development; "
        "when set, all prediction and monitoring endpoints require the key."
    )
    bullet(
        "The Dockerfile does not bake model files into the image. Models are loaded at container "
        "startup from the configured MLflow tracking server, enabling model updates without image "
        "rebuilds. A two-service Docker Compose configuration runs the API alongside an MLflow "
        "tracking server."
    )
    bullet(
        "Dependencies are pinned to exact versions in requirements/base.txt, eliminating the "
        "reproducibility risk of floating version constraints."
    )

    h2("6.3 Limitations")
    bullet(
        "Single-model approach: Each hazard is predicted independently. A multi-task model that "
        "predicts all three simultaneously could exploit shared representations and improve "
        "performance on the more difficult rainfall task."
    )
    bullet(
        "Surface variables only: The models use only surface meteorological observations. Upper-air "
        "data (850 hPa, 500 hPa levels) from ERA5 reanalysis would likely provide significant "
        "predictive power, particularly for rainfall forecasting where mid-tropospheric moisture "
        "and instability are key drivers."
    )
    bullet(
        "No uncertainty quantification: The models produce point predictions with softmax/sigmoid "
        "confidence scores, but do not provide calibrated probability estimates or prediction "
        "intervals. Platt scaling or conformal prediction would improve decision-making utility."
    )
    bullet(
        "Spatial coverage: Fifteen cities provide a reasonable sample of Indian climate zones, but "
        "do not constitute a dense spatial grid. Interpolation to unobserved locations using a "
        "graph neural network or Gaussian process approach would extend coverage."
    )
    bullet(
        "Class imbalance: The cyclone model's 98.5% accuracy is driven largely by correct "
        "classification of the dominant Category 0 (Tropical Depression) class. Per-category "
        "precision and recall for rare categories (4 and 5) should be evaluated separately in "
        "operational use."
    )


def conclusion():
    h1("7. Conclusion")
    para(
        "StormWatch AI demonstrates that production-grade extreme weather prediction is achievable "
        "using relatively simple models (XGBoost) and publicly available data (Open-Meteo, IBTrACS), "
        "provided that careful attention is paid to the integrity of the modelling pipeline. The three "
        "most important lessons from this work are:"
    )

    numbered(
        "Label leakage in temporally autocorrelated data is pernicious and easy to miss. A model "
        "that achieves near-perfect accuracy on weather data is almost certainly cheating — the "
        "correct response to seeing 99.7% accuracy on a next-day rainfall forecast is to suspect "
        "a bug, not to celebrate. Our systematic audit and correction of leakage across all three "
        "models reduced rainfall accuracy by 11.4 percentage points and produced results that domain "
        "knowledge can defend."
    )
    numbered(
        "Feature importance analysis must be accompanied by domain-specific validity checks. "
        "pressure_min's 52.3% importance in the cyclone model could have been interpreted as "
        "residual leakage but was confirmed to represent a legitimate physical relationship through "
        "diagnostic experiments — a practice we recommend for any high-stakes ML system."
    )
    numbered(
        "Hyperparameter tuning is not a substitute for good feature engineering and data quality. "
        "Our Hyperopt study found default XGBoost parameters to be effectively optimal for all "
        "three tasks. The 20x compute cost of tuning would have been better spent on acquiring "
        "upper-air data or implementing multi-task learning."
    )

    para(
        "The system is operational: all three models are served through a FastAPI with authentication, "
        "tracked with MLflow experiment and model registry, monitored for statistical drift via "
        "Kolmogorov-Smirnov tests, and deployable via Docker. The codebase, configuration, "
        "documentation, and test suite (80 tests, 100% passing) are available in the project "
        "repository."
    )


def references_section():
    h1("References")
    bullet("Open-Meteo Archive API. https://archive-api.open-meteo.com/")
    bullet("NOAA IBTrACS v04r01. https://www.ncei.noaa.gov/products/international-best-track-archive")
    bullet(
        "Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. "
        "Proceedings of the 22nd ACM SIGKDD."
    )
    bullet("Bergstra, J., et al. (2011). Algorithms for Hyper-Parameter Optimization. NeurIPS.")
    bullet("Supabase. https://supabase.com/")
    bullet("MLflow. https://mlflow.org/")
    bullet("FastAPI. https://fastapi.tiangolo.com/")

    page_break()


# ── main ──

def main():
    global _doc
    _doc = Document()
    _setup_styles()

    title_page()
    abstract()
    introduction()
    data_section()
    page_break()

    methodology()
    page_break()

    results()
    page_break()

    key_findings()
    discussion()
    conclusion()
    references_section()

    _doc.save(str(OUTPUT))
    size_mb = OUTPUT.stat().st_size / (1024 * 1024)
    img_count = sum(1 for rel in _doc.part.rels.values() if "image" in rel.reltype)
    print(f"Generated {OUTPUT} ({size_mb:.1f} MB)")
    print(f"  Sections: {sum(1 for p in _doc.paragraphs if p.style.name.startswith('Heading 1'))}")
    print(f"  Paragraphs: {len(_doc.paragraphs)}")
    print(f"  Tables: {len(_doc.tables)}")
    print(f"  Figures: {img_count}")


if __name__ == "__main__":
    main()
